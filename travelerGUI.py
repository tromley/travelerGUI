#this is for a dictionary method I used to write lists to a dictionary
from collections import defaultdict
#Handles saving the completed word document
import sys
import os
#PyQt4 is the GUI components
from PyQt4 import QtCore, QtGui
#docx is the name of the plugin that writes the word document
from docx import *
from docx.enum.style import *
from docx.shared import *
from docx.text.parfmt import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.exceptions import PackageNotFoundError

#For error message boxes
import ctypes
import copy



def Mbox(title, text, style):
    """Displays an error message box when the program can't find an empty traveler template

    :param title: A string for the window title
    :param text: A string that is displayed in the body of the message box
    :param style: An integer that corresponds to the style of window that is displayed
    :return: None
    """
    ctypes.windll.user32.MessageBoxW(0, text, title,style)




#Word Document Writing Method Section

arrow_code = '➔ ' #The arrow that is used in the traveler document template here so I don't have to type it every time
#This list is referred to in the "Master_Writer" function below. These "support_keys" correspond to the 1st item(0 index) in the
#row_contents variable defined in the Buildmeatraveler Class defined below. They are used to determine what type of row the master_writer
# function needs to write.
support_keys = ['ROUTING STEP','1ST LEVEL','2ND LEVEL','3RD LEVEL','TECH ENTRY LINE']

def master_writer(filepath, filedirectory, document_name, new_template_location):
    """Reads the row_contents dictionary of class UI_Buildmeatraveler and calls the appropriate docx writing function
    as well as the appropriate support text writing function. It defines the word styles in the new traveler document as
    well the margins. It then adds a rework table and runs saves the traveler function.

    :param filepath: A string of the filepath of where the traveler is being saved
    :param filedirectory: A string of the filepath of where the traveler is being saved without the name of the traveler document
    :param document_name: A string of the name of the traveler document that the user saved
    :param new_template_location: A string that if the "Set Empty Traveler" Button hasn't been pressed this is "". If the "Set Empty Traveler"
    button has been pressed this is the new filepath of the empty traveler template
    :return: None
    """


    if new_template_location == "": #Looks for the default empty barcode template and initializes a document object
        try:
            barcode_template_location = '//vcafps01/Projects-2/Manufacturing Engineering/Build a Traveler/Empty Traveler Template.docx'
            traveler = Document(barcode_template_location)
        except:
            Mbox('Can\'t Find Empty Template',
                 'Please press the "Set Empty Template Location" button before saving this traveler.', 0)
            traveler = Document() #a new document needs to be made even if the empty template is not found. This prevents other errors
            pass
    else:
        traveler = Document(new_template_location)
    sections = traveler.sections
    section = sections[0]
    section.top_margin = Inches(.25)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
    section.bottom_margin = Inches(.2)
    DefineStyles('Routing Step', 'Arial', Pt(9), True, traveler)
    DefineStyles('Instructions', 'Arial', Pt(9), False, traveler)
    DefineStyles('Sub Instructions', 'Arial', Pt(9), False, traveler)
    DefineStyles('Note', 'Arial', Pt(9), True, traveler)
    DefineStyles('Routing Step Number', 'Arial', Pt(12), True, traveler)
    DefineStyles('Rework Table', 'Arial', Pt(11), True, traveler)
    DefineStyles('Rework Table Heading', 'Arial', Pt(12), True, traveler)
    try: #Looks for an existing support text document and deletes it to prevent weird overwrite errors
        os.remove(filepath)
        support_text = open(filedirectory + document_name, 'w')
    except:
        support_text = open(filedirectory + document_name, 'w')

    row_contents = ui.row_contents

    for row in row_contents:
        specefic_row = row_contents[row]
        type_of_row = specefic_row[-1] #the last item in each row is a string that identifies what type of row is being added
        if type_of_row == ui.keyRoutingStep:
            add_RoutingStepDescription(specefic_row, traveler)
            support_add_RoutingStep(support_keys[0],support_text,specefic_row)
        elif type_of_row == ui.keyFirstLevelInstruction:
            add_FirstLevelInstruction(specefic_row, traveler)
            support_add_LowerLevelInstructions(support_keys[1],support_text,specefic_row)
        elif type_of_row == ui.keySecondLevelInstruction:
            add_SecondLevelInstruction(specefic_row, traveler)
            support_add_LowerLevelInstructions(support_keys[2],support_text, specefic_row)
        elif type_of_row == ui.keyAdditionalTechEntry:
            add_TechEntryField(specefic_row,traveler)
            support_add_TechEntryField(support_keys[4],support_text,specefic_row)
        elif type_of_row == ui.keyThirdLevelInstruction:
            add_ThirdLevelInstruction(specefic_row, traveler)
            support_add_LowerLevelInstructions(support_keys[3],support_text, specefic_row)
        else:
            print("The row_contents identifier key does not match any of the available row types.")

    add_ReworkTable(traveler)
    save_traveler(filepath,traveler)
    support_text.close()

#.txt support file writing section
def support_add_TechEntryField(support_key,support_text,specefic_row):
    """Adds a line to the support text document that represents a Tech Entry Dropdown Box

    :param support_key: A string that corresponds to the support_keys list that identifies what type of line is written
    :param support_text: A string representing the filepath of the support text document. Text document name included
    :param specefic_row: A list of all the Qt Objects in the row that is being written. The last item is a string identifier.
    The indexes of the list are as follows: 0 - TechEntryDropdown Box 1- Label
    :return: None
    """
    #key identifier
    support_text.write(support_key+'~')
    # combobox text
    support_text.write(specefic_row[0].currentText() + '~')
    # label text
    support_text.write(specefic_row[1].text() + '\n')

def support_add_LowerLevelInstructions(support_key,support_text,specefic_row):
    """Adds a line to the support text document that represents a 1st 2nd or 3rd level instruction

    :param support_key: A string that corresponds to the support_keys list that identifies what type of line is written
    :param support_text: A string representing the filepath of the support text document. Text document name included
    :param specefic_row: A list of all the Qt Objects in the row that is being written. The last item is a string identifier.
    The indexes of the list are as follows: 0 - TechEntryDropdown Box 1- Label 2 - Instruction Description
    :return: None
    """


    support_text.write(support_key+'~')     #key identifier

    support_text.write(specefic_row[0].currentText() + '~') # combobox text

    support_text.write(specefic_row[1].text() + '~') # label text

    support_text.write(specefic_row[2].text() + '\n') # description text

def support_add_RoutingStep(support_key,support_text,specefic_row):
    """Adds a line to the support text document that represents a Routing Step

    :param support_key: A string that corresponds to the support_keys list that identifies what type of line is written
    :param support_text: A string representing the filepath of the support text document. Text document name included
    :param specefic_row: A list of all the Qt Objects in the row that is being written. The last item is a string identifier.
    The indexes of the list are as follows: 0 - TechEntryDropdown Box 1- Label 2- Routing Step Number 3 - Routing Step Description
    :return: None
    """
    #key identifier
    support_text.write(support_key+'~')
    #combobox text
    support_text.write(specefic_row[0].currentText()+'~')
    #label text
    support_text.write(specefic_row[1].text()+'~')
    #Routing Step text
    support_text.write(specefic_row[2].text()+'~')
    #description text
    support_text.write(specefic_row[3].text()+'\n')
#end .txt support file writing section

def add_RoutingStepDescription(row, traveler):
    """ Adds everything a routing step except the number to the word document. Also defines the line spacing of a Routing Step

    :param row: a list containing all the Qt objects in a specefic row. THe last item is a string identifier
    :param traveler: A Document object that represents the word document that is being written
    :return: None
    """
    tech_entry_box = row[0]
    row_label = row[1]
    routing_step_number = row[2]
    routing_step_description = row[3]
    working_paragraph = add_RoutingStepNumber(routing_step_number,traveler)
    paragraph_format = working_paragraph.paragraph_format
    paragraph_format.space_after = Pt(6)
    new_run = working_paragraph.add_run(routing_step_description.text()) #the routing step description is added as a run
                                                                        # to the paragraph that contains just the routing step number
    if new_run != None:
        font = new_run.font
        font.name = 'Arial'
        font.size = Pt(9)
        font.bold = True
    last_run = tech_entry_input(tech_entry_box,working_paragraph)
    if last_run != None:
        font = last_run.font
        font.name = 'Arial'
        font.size = Pt(9)
        font.bold = False

def add_FirstLevelInstruction(row, traveler):
    """ Adds a first level instruction to the word document. Also defines the line spacing of a first level instruction

    :param row: a list containing all the Qt objects in a specefic row. THe last item is a string identifier
    :param traveler: A Document object that represents the word document that is being written
    :return: None
    """
    tech_entry_box = row[0]
    row_label = row[1]
    instruction_description = row[2]
    addition = traveler.add_paragraph(arrow_code + instruction_description.text())
    addition.style = 'Instructions'
    paragraph_format = addition.paragraph_format
    paragraph_format.space_after = Pt(6)
    add_Tabstop(addition)
    tech_entry_input(tech_entry_box, addition)

def add_SecondLevelInstruction(row, traveler):
    """ Adds a 2nd level instruction to the word document. Also defines the line spacing of a 2nd level instruction

    :param row: a list containing all the Qt objects in a specefic row. THe last item is a string identifier
    :param traveler: A Document object that represents the word document that is being written
    :return: None
    """
    tech_entry_box = row[0]
    row_label = row[1]
    sub_instruction_description = row[2]
    addition = traveler.add_paragraph("       "+arrow_code+sub_instruction_description.text())
    add_Tabstop(addition)
    paragraph_format = addition.paragraph_format
    paragraph_format.space_after = Pt(6)
    addition.style = 'Sub Instructions'
    tech_entry_input(tech_entry_box,addition)

def add_TechEntryField(row,traveler):
    """ Adds a Tech Entry Field to the word document.

    :param row: a list containing all the Qt objects in a specefic row. THe last item is a string identifier
    :param traveler: A Document object that represents the word document that is being written
    :return: None
    """
    tech_entry_box = row[0]
    addition = traveler.add_paragraph("")
    add_Tabstop(addition)
    addition.style = 'Instructions'
    tech_entry_input(tech_entry_box,addition)

def add_ThirdLevelInstruction(row, traveler):
    """ Adds a 3rd level instruction to the word document. Also defines the line spacing of a third level instruction

    :param row: a list containing all the Qt objects in a specefic row. THe last item is a string identifier
    :param traveler: A Document object that represents the word document that is being written
    :return: None
    """
    tech_entry_box = row[0]
    row_label = row[1]
    thirdLevelInstructionDescription = row[2]
    addition = traveler.add_paragraph("            "+arrow_code+thirdLevelInstructionDescription.text())
    add_Tabstop(addition)
    paragraph_format = addition.paragraph_format
    paragraph_format.space_after = Pt(6)
    addition.style = 'Sub Instructions'
    tech_entry_input(tech_entry_box,addition)

def add_RoutingStepNumber(routing_step_number, traveler):
    """ Adds the routing step number to the word document

    :param routing_step_number: A string containing the routing step number
    :param traveler: A document object that is being written
    :return: Paragraph object that contains the routing step number
    """
    addition = traveler.add_paragraph("(" + routing_step_number.text() + ") ")
    add_Tabstop(addition)
    addition.style = 'Routing Step Number'
    return addition

def DefineStyles(stylename,fontname,fontsize,bold,traveler): #inherits normal styles of a blank word document
    """Add's the specefic styles needed for the different types of steps to the word document

    :param stylename: A string containing the name of the style to be added
    :param fontname: A string containing the name of the font to be used in the style
    :param fontsize: An integer size that defines the size of the font. A Pt(integer) conversion is used
    :param bold: A boolean that determines if the style will have bolded text
    :param traveler: A document object to add the style to
    :return: None
    """
    styles = traveler.styles
    if stylename == 'Routing Step Number' or 'Instructions':
        style = styles.add_style(stylename, WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles.add_style(stylename, WD_STYLE_TYPE.CHARACTER)
    style.base_style = styles['Normal']
    font = style.font
    font.name = fontname
    font.size = fontsize
    font.bold = bold


def add_Tabstop(paragraph):
    """Add's a tabstop to word document line being written

    :param paragraph: A paragraph object that the tabstop is added to
    :return: None
    """
    paragraph_format = paragraph.paragraph_format
    tab_stops = paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Inches(4.5))

def tech_entry_input(tech_entry_combobox,working_paragraph):
    """Adds the operator input lines to the end of the current line in the word document

    :param tech_entry_combobox:  A Qt Combobox Object that is used as the Tech Entry Dropdown
    :param working_paragraph: A paragraph object that is currently being written
    :return: None if the Combobox's current value is "None". Else returns a paragraph object with tech entry field added
    """
    tech_entry_type = tech_entry_combobox.currentText()
    if tech_entry_type == "None":
        return None
    elif tech_entry_type == "SN/Rev":
        last_run =working_paragraph.add_run("\t"+"SN: ________________   REV: ________________")
        paragraph_format = working_paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)
    elif tech_entry_type == "Rev":
        last_run =working_paragraph.add_run("\t"+"REV: ______________")
        paragraph_format = working_paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)
    elif tech_entry_type == "ID":
        last_run =working_paragraph.add_run("\t"+"Tool ID: ______________")
        paragraph_format = working_paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)
    elif tech_entry_type == "Operator/Date":
        last_run =working_paragraph.add_run("\t"+"Operator: ________________    Date: ____________")
        paragraph_format = working_paragraph.paragraph_format
        paragraph_format.space_before = Pt(18)
        paragraph_format.space_after = Pt(6)
    if tech_entry_type != "Operator/Date":
        paragraph_format = working_paragraph.paragraph_format
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)
    return last_run

def add_ReworkTable(traveler):
    """Add's a rework table to the end of the word document

    :param traveler: A document object that is being written
    :return: None
    """
    tables = traveler.tables
    if len(tables) >= 2:
        return None
    else:
        traveler.add_page_break()
        title_of_table = traveler.add_paragraph('Production Rework Record')
        title_of_table.style = 'Rework Table Heading'
        title_of_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        reworkTable = traveler.add_table(rows = 25, cols = 8)
        headings = ['Ref.\nOper','Ref.\nECO','NMR/\nFailure\nReport','Description','Date','Test\nStamp','Oper.\nStamp','Q.A.\nStamp']
        for i in range(0,8):
            cell = reworkTable.cell(0,i)
            cell.text = headings[i]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.style = 'Rework Table'
            if i != 3:
                cell.width = Inches(.7)
        cell = reworkTable.cell(0,3)
        cell.width = Inches(2.8)
        styles = traveler.styles
        for i in range(0,25): #This code was found online. As of this writing there is no convenient way to resize table row heights
            row = reworkTable.rows[i]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '404.7132857') #I'm not sure what units the 404.7132857 number is in but it is equivalent to .28 inches.
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
        reworkTable.alignment = WD_TABLE_ALIGNMENT.CENTER
        reworkTable.style = 'Table Grid'

def save_traveler(filepath,traveler):
    """Saves the traveler word document to a specefied location

    :param filepath:  A string of the filepath of where the traveler is being saved
    :param traveler: A document object that is going to be saved
    :return: None
    """
    try:
        traveler.save(filepath)
        os.startfile(filepath) #Runs the traveler immediately after saving
    except:
        ctypes.windll.user32.MessageBoxW(0, 'Error saving file! Make Sure that the Word Document you are saving to is not open!', 'Save Error', 0)
#End Word Document Writing Section


#loading text document section

def load_support_text_file(filepath,document_name):
    """Similar to the master writer function. Reads each line in the support text document and calls on the appropriate
    function to add a line to the GUI

    :param filepath: A string containing the filepath of support text document to be opened
    :param document_name: A string ending in .txt contianing the name of the support text file
    :return: None
    """
    support_text = open(filepath,'r')
    full_name = document_name
    list_of_parts = full_name.split('.')
    try: #automatically fills in the PN and Rev # of the traveler if the support text file is named in the form TR-PN_Rev.txt
        pure_travler_name= list_of_parts[0]
        pn_and_rev = pure_travler_name.split('-')
        list_of_pn_and_rev = pn_and_rev[1].split('_')
        ui.entryPartNumber.setText(list_of_pn_and_rev[0])
        ui.entryRev.setText(list_of_pn_and_rev[1])
    except IndexError: #Raises if the support text file is not named in the form TR-PN_Rev.txt
        print("Couldn't find Rev Number. For automatic entry of PN and Rev # in the travelerGUI please name your document"
              " in the following format: TR-PN_Rev")
    for line in support_text:
        individual_entries = line.split('~')
        stripped_last_entry = individual_entries[-1].rstrip('\n') #Removes the new line character after at the end of every line
        individual_entries.pop()
        individual_entries.append(stripped_last_entry) #Adds a '' string to the end of the list
        type_of_row = individual_entries[0].upper()
        if type_of_row == support_keys[0]:         #routing step
            ui.from_load_addRoutingStep(individual_entries[1],individual_entries[2],individual_entries[3],individual_entries[4])
        elif type_of_row == support_keys[1]:         #1st level
            ui.from_load_add1stLevelInstruction(individual_entries[1],individual_entries[2],individual_entries[3])
        elif type_of_row == support_keys[2]:         #second level
            ui.from_load_add2ndLevelInstruction(individual_entries[1],individual_entries[2],individual_entries[3])
        elif type_of_row == support_keys[3]:         #3rd level
            ui.from_load_add3rdLevelInstruction(individual_entries[1], individual_entries[2], individual_entries[3])
        elif type_of_row == support_keys[4]:         #tech entry field
            ui.from_load_addAdditionalTechEntryRow(individual_entries[1],individual_entries[2])
        else:
            print("Error in row: "+line)



#UI Interface and dynamic widget adding methods
#Most of the code below was obtained by converting the Qt Creator code to python. All comments are written by me.
try:
    _fromUtf8 = QtCore.QString.fromUtf8
except AttributeError:
    def _fromUtf8(s):
        return s

try:
    _encoding = QtGui.QApplication.UnicodeUTF8
    def _translate(context, text, disambig):
        return QtGui.QApplication.translate(context, text, disambig, _encoding)

except AttributeError:
    def _translate(context, text, disambig):
        return QtGui.QApplication.translate(context, text, disambig)

class Ui_BuildMeATraveler(object):
    new_empty_template_location = "" #Only changes to a new filepath when the "Set Empty Traveler" Button is pressed
    row_contents = {} #dictionary that is appended when a new row of qt objects are added through any of the "Add XXXXX" buttons
                      # Keys represent row numbers, Each key stores a list of Qt Objects and an identifier that is read by the master_writer function
    row_contents = defaultdict(list)
    def setupUi(self, BuildMeATraveler): #initializes the UI
        BuildMeATraveler.setObjectName(_fromUtf8("BuildMeATraveler"))
        BuildMeATraveler.resize(1077, 638)
        self.centralWidget = QtGui.QWidget(BuildMeATraveler)
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Expanding, QtGui.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.centralWidget.sizePolicy().hasHeightForWidth())
        self.centralWidget.setSizePolicy(sizePolicy)
        self.centralWidget.setObjectName(_fromUtf8("centralWidget"))
        self.gridLayout_2 = QtGui.QGridLayout(self.centralWidget)
        self.gridLayout_2.setMargin(11)
        self.gridLayout_2.setSpacing(6)
        self.gridLayout_2.setObjectName(_fromUtf8("gridLayout_2"))
        self.horizontalLayout_2 = QtGui.QHBoxLayout()
        self.horizontalLayout_2.setMargin(11)
        self.horizontalLayout_2.setSpacing(6)
        self.horizontalLayout_2.setObjectName(_fromUtf8("horizontalLayout_2"))
        self.bottomframe = QtGui.QHBoxLayout()
        self.bottomframe.setContentsMargins(11, 11, 0, 11)
        self.bottomframe.setSpacing(500)
        self.bottomframe.setObjectName(_fromUtf8("bottomframe"))
        self.btnExitTraveler = QtGui.QPushButton(self.centralWidget)
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.btnExitTraveler.sizePolicy().hasHeightForWidth())
        self.btnExitTraveler.setSizePolicy(sizePolicy)
        self.btnExitTraveler.setObjectName(_fromUtf8("btnExitTraveler"))
        self.bottomframe.addWidget(self.btnExitTraveler)
        self.btnSaveTraveler = QtGui.QPushButton(self.centralWidget)
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.btnSaveTraveler.sizePolicy().hasHeightForWidth())
        self.btnSaveTraveler.setSizePolicy(sizePolicy)
        self.btnSaveTraveler.setObjectName(_fromUtf8("btnSaveTraveler"))
        self.bottomframe.addWidget(self.btnSaveTraveler)
        self.horizontalLayout_2.addLayout(self.bottomframe)
        self.gridLayout_2.addLayout(self.horizontalLayout_2, 2, 0, 1, 1)
        self.verticalLayout_2 = QtGui.QVBoxLayout()
        self.verticalLayout_2.setMargin(11)
        self.verticalLayout_2.setSpacing(6)
        self.verticalLayout_2.setObjectName(_fromUtf8("verticalLayout_2"))
        self.addLineFrame = QtGui.QFrame(self.centralWidget)
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Preferred, QtGui.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.addLineFrame.sizePolicy().hasHeightForWidth())
        self.addLineFrame.setSizePolicy(sizePolicy)
        self.addLineFrame.setMinimumSize(QtCore.QSize(1, 175))
        self.addLineFrame.setFrameShape(QtGui.QFrame.NoFrame)
        self.addLineFrame.setFrameShadow(QtGui.QFrame.Sunken)
        self.addLineFrame.setObjectName(_fromUtf8("addLineFrame"))
        self.btnAddRoutingStep = QtGui.QPushButton(self.addLineFrame)
        self.btnAddRoutingStep.setGeometry(QtCore.QRect(10, 70, 131, 28))
        self.btnAddRoutingStep.setObjectName(_fromUtf8("btnAddRoutingStep"))
        self.btnAddRoutingInstruction = QtGui.QPushButton(self.addLineFrame)
        self.btnAddRoutingInstruction.setGeometry(QtCore.QRect(160, 70, 171, 28))
        self.btnAddRoutingInstruction.setAutoExclusive(False)
        self.btnAddRoutingInstruction.setObjectName(_fromUtf8("btnAddRoutingInstruction"))
        self.btnAddSubInstruction = QtGui.QPushButton(self.addLineFrame)
        self.btnAddSubInstruction.setGeometry(QtCore.QRect(350, 70, 171, 28))
        self.btnAddSubInstruction.setObjectName(_fromUtf8("btnAddSubInstruction"))
        self.btnAdd3rdLevelInstruction = QtGui.QPushButton(self.addLineFrame)
        self.btnAdd3rdLevelInstruction.setGeometry(QtCore.QRect(730, 70, 181, 28))
        self.btnAdd3rdLevelInstruction.setObjectName(_fromUtf8("btnAdd3rdLevelInstruction"))
        self.entryPartNumber = QtGui.QLineEdit(self.addLineFrame)
        self.entryPartNumber.setGeometry(QtCore.QRect(120, 10, 201, 22))
        self.entryPartNumber.setText(_fromUtf8(""))
        self.entryPartNumber.setObjectName(_fromUtf8("entryPartNumber"))
        self.labelPartNumber = QtGui.QLabel(self.addLineFrame)
        self.labelPartNumber.setGeometry(QtCore.QRect(20, 10, 91, 21))
        self.labelPartNumber.setObjectName(_fromUtf8("labelPartNumber"))
        self.entryPartDescription = QtGui.QLineEdit(self.addLineFrame)
        self.entryPartDescription.setGeometry(QtCore.QRect(120, 40, 201, 22))
        self.entryPartDescription.setReadOnly(True)
        self.entryPartDescription.setObjectName(_fromUtf8("entryPartDescription"))
        self.labelPartDescription = QtGui.QLabel(self.addLineFrame)
        self.labelPartDescription.setGeometry(QtCore.QRect(20, 40, 91, 21))
        self.labelPartDescription.setObjectName(_fromUtf8("labelPartDescription"))
        self.btnRemovePreviousEntry = QtGui.QPushButton(self.addLineFrame)
        self.btnRemovePreviousEntry.setGeometry(QtCore.QRect(10, 130, 171, 28))
        self.btnRemovePreviousEntry.setObjectName(_fromUtf8("btnRemovePreviousEntry"))
        self.labelRev = QtGui.QLabel(self.addLineFrame)
        self.labelRev.setGeometry(QtCore.QRect(340, 10, 81, 21))
        self.labelRev.setObjectName(_fromUtf8("labelRev"))
        self.entryRev = QtGui.QLineEdit(self.addLineFrame)
        self.entryRev.setGeometry(QtCore.QRect(430, 10, 201, 22))
        self.entryRev.setText(_fromUtf8(""))
        self.entryRev.setObjectName(_fromUtf8("entryRev"))
        self.btnAddTechEntryLine = QtGui.QPushButton(self.addLineFrame)
        self.btnAddTechEntryLine.setGeometry(QtCore.QRect(540, 70, 171, 28))
        self.btnAddTechEntryLine.setObjectName(_fromUtf8("btnAddTechEntryLine"))
        self.btnSetTemplateLocation = QtGui.QPushButton(self.addLineFrame)
        self.btnSetTemplateLocation.setGeometry(QtCore.QRect(710, 10, 201, 28))
        self.btnSetTemplateLocation.setFocusPolicy(QtCore.Qt.NoFocus)
        self.btnSetTemplateLocation.setObjectName(_fromUtf8("btnSetTemplateLocation"))
        self.dropRoutingStepIncrement = QtGui.QComboBox(self.addLineFrame)
        self.dropRoutingStepIncrement.setGeometry(QtCore.QRect(370, 130, 73, 21))
        self.dropRoutingStepIncrement.setObjectName(_fromUtf8("dropRoutingStepIncrement"))
        self.dropRoutingStepIncrement.addItem(_fromUtf8(""))
        self.dropRoutingStepIncrement.addItem(_fromUtf8(""))
        self.btnPopulateRoutings = QtGui.QPushButton(self.addLineFrame)
        self.btnPopulateRoutings.setGeometry(QtCore.QRect(230, 130, 131, 28))
        self.btnPopulateRoutings.setObjectName(_fromUtf8("btnPopulateRoutings"))
        self.btnClearRoutings = QtGui.QPushButton(self.addLineFrame)
        self.btnClearRoutings.setGeometry(QtCore.QRect(450, 130, 93, 28))
        self.btnClearRoutings.setObjectName(_fromUtf8("btnClearRoutings"))
        self.pushButton = QtGui.QPushButton(self.addLineFrame)
        self.pushButton.setGeometry(QtCore.QRect(640, 130, 181, 28))
        self.pushButton.setObjectName(_fromUtf8("pushButton"))
        self.verticalLayout_2.addWidget(self.addLineFrame)
        self.gridLayout_2.addLayout(self.verticalLayout_2, 0, 0, 1, 1)
        self.scrollArea = QtGui.QScrollArea(self.centralWidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName(_fromUtf8("scrollArea"))
        self.addInfoWidget = QtGui.QWidget()
        self.addInfoWidget.setGeometry(QtCore.QRect(0, 0, 1053, 326))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Expanding, QtGui.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.addInfoWidget.sizePolicy().hasHeightForWidth())
        self.addInfoWidget.setSizePolicy(sizePolicy)
        self.addInfoWidget.setObjectName(_fromUtf8("addInfoWidget"))
        self.Addinfogrid = QtGui.QGridLayout(self.addInfoWidget)
        self.Addinfogrid.setMargin(11)
        self.Addinfogrid.setHorizontalSpacing(6)
        self.Addinfogrid.setVerticalSpacing(3)
        self.Addinfogrid.setObjectName(_fromUtf8("Addinfogrid"))
        self.scrollArea.setWidget(self.addInfoWidget)
        self.gridLayout_2.addWidget(self.scrollArea, 1, 0, 1, 1)
        BuildMeATraveler.setCentralWidget(self.centralWidget)
        self.menuBar = QtGui.QMenuBar(BuildMeATraveler)
        self.menuBar.setGeometry(QtCore.QRect(0, 0, 1077, 26))
        self.menuBar.setObjectName(_fromUtf8("menuBar"))
        self.menuBuild_A_Traveler = QtGui.QMenu(self.menuBar)
        self.menuBuild_A_Traveler.setObjectName(_fromUtf8("menuBuild_A_Traveler"))
        BuildMeATraveler.setMenuBar(self.menuBar)
        self.mainToolBar = QtGui.QToolBar(BuildMeATraveler)
        self.mainToolBar.setObjectName(_fromUtf8("mainToolBar"))
        BuildMeATraveler.addToolBar(QtCore.Qt.TopToolBarArea, self.mainToolBar)
        self.statusBar = QtGui.QStatusBar(BuildMeATraveler)
        self.statusBar.setObjectName(_fromUtf8("statusBar"))
        BuildMeATraveler.setStatusBar(self.statusBar)
        self.actionBuildFromTextFile = QtGui.QAction(BuildMeATraveler)
        self.actionBuildFromTextFile.setObjectName(_fromUtf8("actionBuildFromTextFile"))
        self.menuBuild_A_Traveler.addAction(self.actionBuildFromTextFile)
        self.menuBar.addAction(self.menuBuild_A_Traveler.menuAction())

        self.retranslateUi(BuildMeATraveler)
        QtCore.QMetaObject.connectSlotsByName(BuildMeATraveler)

        # connect statements trigger specefic functions to run when buttons are clicked
        self.actionBuildFromTextFile.triggered.connect(self.support_text_open)
        self.btnClearRoutings.clicked.connect(self.clearRoutings)
        self.btnPopulateRoutings.clicked.connect(self.populateRoutings)
        self.btnAdd3rdLevelInstruction.clicked.connect(self.add3rdLevelInstruction)
        self.btnExitTraveler.clicked.connect(self.close_application)
        self.btnSaveTraveler.clicked.connect(self.file_save)
        self.btnSetTemplateLocation.clicked.connect(empty_template_open)
        self.btnAddRoutingStep.clicked.connect(self.addRoutingStep)
        self.btnAddRoutingInstruction.clicked.connect(self.add1stLevelInstruction)
        self.btnAddSubInstruction.clicked.connect(self.add2ndLevelInstruction)
        self.btnRemovePreviousEntry.clicked.connect(self.remove_previous_entry)
        self.btnAddTechEntryLine.clicked.connect(self.addAdditionalTechEntryRow)
        # variables used to keep track of where to add rows in the Addinfogrid
        self.routing_step_counter = 0
        self.current_row = 1
        self.naming_counter = 1
        # keys are read by python -docx method to determine what line to write
        self.keyNone = "None"
        self.keyRoutingStep = self.btnAddRoutingStep.text()
        self.keyFirstLevelInstruction = self.btnAddRoutingInstruction.text()
        self.keySecondLevelInstruction = self.btnAddSubInstruction.text()
        self.keyNoRoutingstep = "No Routing Step"
        self.keyAdditionalTechEntry = self.btnAddTechEntryLine.text()
        self.keyThirdLevelInstruction = self.btnAdd3rdLevelInstruction.text()

    def support_text_open(self):
        """Called when the load from a .txt file action item is pressed. Resets the class variables and deletes any rows
        that might of already been added to the TravelerGUI. Then calls on the load_support_text_file function
        :return: None
        """
        filepath = QtGui.QFileDialog.getOpenFileName(None,'Open Support Text File')
        split_path = filepath.rsplit('/', maxsplit=1)
        document_name = split_path[-1] + '.txt'
        self.routing_step_counter = 0
        self.current_row = 1
        self.naming_counter = 1
        while self.Addinfogrid.count():
            child = self.Addinfogrid.takeAt(0)
            child.widget().deleteLater()
        self.row_contents.clear()
        load_support_text_file(filepath,document_name)


    def file_save(self):
        """Opens a save dialog box for user to choose where to save traveler. Adds a .docx extension to the file. Runs Master_writer function

        :return: None
        """
        try:
            filepath = QtGui.QFileDialog.getSaveFileName(None,'Save File','TR-'+self.entryPartNumber.text()+'_'+self.entryRev.text())
            split_path = filepath.rsplit('/', maxsplit=1)
            document_name = split_path[-1] +'.txt'
            document_directory = split_path[-2] + '/'
            if filepath[-1] != 'x' or filepath[-2] != 'c' or filepath[-3] !='o' or filepath[-4] != 'd' or filepath[-5] != '.':
                filepath = filepath +'.docx'
            master_writer(filepath, document_directory, document_name, self.new_empty_template_location)
        except IndexError:
            None
    def clearRoutings(self):
        """Clears all routing step numbers in the travelerGUI

        :return: None
        """
        for row in ui.row_contents:
            specefic_row = ui.row_contents[row]
            type_of_row = specefic_row[-1]
            if type_of_row =='Add Routing Step':
                specefic_row[2].setText("")
                self.routing_step_counter = 0
            else:
                None

    def populateRoutings(self):
        """Populates all available Routing Step LineEdits in increments determined by dropRoutingStepIncrement Combobox

        :return: None
        """
        for row in ui.row_contents:
            specefic_row = ui.row_contents[row]
            type_of_row = specefic_row[-1]
            if type_of_row == self.keyRoutingStep:
             if self.dropRoutingStepIncrement.currentText() =='10':
                self.routing_step_counter += 10
                specefic_row[2].setText(str(self.routing_step_counter))
             elif self.dropRoutingStepIncrement.currentText() =='5':
                self.routing_step_counter +=5
                specefic_row[2].setText(str(self.routing_step_counter))
             else:
                ctypes.windll.user32.MessageBoxW(0, 'Please select an increment.', 'Select an increment', 0)

    def remove_previous_entry(self):
        """ Removes all entries in Routing Step LineEdits

        :return: None
        """
        try:
            self.routing_step_counter = 0
            previous_entry_row = self.current_row - 1
            previous_entry_contents = self.row_contents.pop(previous_entry_row)
            for item in previous_entry_contents:
                if type(item) == str:
                    previous_entry_contents.remove(item)
                else:
                    item.deleteLater()
                    item = None
            if self.current_row > 1:
                self.current_row = self.current_row - 1
        except KeyError:
            print("No more rows left to remove!")
    def add3rdLevelInstruction(self):
        """ Adds a Label, Instruction LineEdit and Tech Entry Combobox to the UI. Updates Row_Contents, current_row and Naming_counter

        :return: None
        """
        # create label
        self.label = QtGui.QLabel(self.centralWidget)
        self.label.setText(self.btnAdd3rdLevelInstruction.text())
        self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        self.label.setSizePolicy(sizePolicy)
        self.Addinfogrid.addWidget(self.label, self.current_row, 1, 1, 1)
        # create Description box
        self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
        self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
        self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
        # create entry combo
        self.setEntryBoxSelections(self.current_row)
        # update variables #also adds keyidentifier described above.
        self.row_contents[self.current_row].append(self.label)
        self.row_contents[self.current_row].append(self.routing_description_line_edit)
        self.row_contents[self.current_row].append(self.keyThirdLevelInstruction)
        self.current_row += 1
        self.naming_counter += 1

    def addAdditionalTechEntryRow(self):
        """ Adds a Label and a combobox to the UI. Updates the row_contents, current_row and naming_counter

        :return: None
        """
        # create label
        self.label = QtGui.QLabel(self.centralWidget)
        self.label.setText(self.btnAddTechEntryLine.text())
        self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        self.label.setSizePolicy(sizePolicy)
        self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
        #create combobox
        combobox = QtGui.QComboBox(self.centralWidget)
        combobox.setObjectName(_fromUtf8("combobox_" + str(self.naming_counter)))
        combobox.addItem(_fromUtf8("Operator/Date"))
        combobox.addItem(_fromUtf8("SN/Rev"))
        combobox.addItem(_fromUtf8("Rev"))
        combobox.addItem(_fromUtf8("ID"))
        self.Addinfogrid.addWidget(combobox, self.current_row, 3, 1, 1)
        #update
        self.row_contents[self.current_row].append(combobox)
        self.row_contents[self.current_row].append(self.label)
        self.row_contents[self.current_row].append(self.keyAdditionalTechEntry)
        self.current_row += 1
        self.naming_counter += 1

    def addRoutingStep(self):
        """Adds a Label, Routing Step LineEdit, Routing Description Line Edit and Combobox to the UI. Updates row_contents
        naming_counter and current_row

        :return: None
        """
        self.routing_step_counter = 0         #reset the routing step numbers so it autopopulates correctly
        # create label
        self.label = QtGui.QLabel(self.centralWidget)
        self.label.setText(self.btnAddRoutingStep.text())
        self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        self.label.setSizePolicy(sizePolicy)
        self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
        # create Routing Step Box
        self.routing_step_line_edit = QtGui.QLineEdit(self.centralWidget)
        self.routing_step_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
        self.Addinfogrid.addWidget(self.routing_step_line_edit, self.current_row, 1, 1, 1)
        self.routing_step_line_edit.setFixedWidth(30)
        # create Descr Box
        self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
        self.routing_description_line_edit.setObjectName(
            _fromUtf8("RoutingDescription_" + str(self.naming_counter)))
        self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
        # create entry combo
        self.setEntryBoxSelections(self.current_row)
        # update variables
        self.row_contents[self.current_row].append(self.label)
        self.row_contents[self.current_row].append(self.routing_step_line_edit)
        self.row_contents[self.current_row].append(self.routing_description_line_edit)
        self.row_contents[self.current_row].append(self.keyRoutingStep)
        self.current_row += 1
        self.naming_counter += 1

    def add1stLevelInstruction(self):
        """ Adds a Label, Instruction Description LineEdit, Entry Combobox to the UI. Updates UI Variables

        :return: None
        """
        # create label
        self.label = QtGui.QLabel(self.centralWidget)
        self.label.setText(self.btnAddRoutingInstruction.text())
        self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        self.label.setSizePolicy(sizePolicy)
        self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
        # create Description box
        self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
        self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
        self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
        # create entry combo
        self.setEntryBoxSelections(self.current_row)
        # update variables #also adds keyidentifier described above.
        self.row_contents[self.current_row].append(self.label)
        self.row_contents[self.current_row].append(self.routing_description_line_edit)
        self.row_contents[self.current_row].append(self.keyFirstLevelInstruction)
        self.current_row += 1
        self.naming_counter += 1

    def add2ndLevelInstruction(self):
        """Adds a label, line_edit , combobox and updates the ui variables. Updates UI Variables

        :return: None
        """
        # create label
        self.label = QtGui.QLabel(self.centralWidget)
        self.label.setText(self.btnAddSubInstruction.text())
        self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
        sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
        self.label.setSizePolicy(sizePolicy)
        self.Addinfogrid.addWidget(self.label, self.current_row, 1, 1, 1)
        # create Description box
        self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
        self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
        self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
        # create entry combo
        self.setEntryBoxSelections(self.current_row)
        # update variables #also adds keyidentifier described above.
        self.row_contents[self.current_row].append(self.label)
        self.row_contents[self.current_row].append(self.routing_description_line_edit)
        self.row_contents[self.current_row].append(self.keySecondLevelInstruction)
        self.current_row += 1
        self.naming_counter += 1

    def close_application(self):
        """

        :return:
        """
        result = QtGui.QMessageBox.question(self.btnExitTraveler, 'Exit Application',
                                            "Are you sure you want to exit?",
                                            QtGui.QMessageBox.Yes, QtGui.QMessageBox.No)
        if result == QtGui.QMessageBox.Yes:
            sys.exit()

    def from_load_add3rdLevelInstruction(self,combobox_text,label_text,description_text):
            # create label
            self.label = QtGui.QLabel(self.centralWidget)
            self.label.setText(label_text)
            self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
            sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
            self.label.setSizePolicy(sizePolicy)
            self.Addinfogrid.addWidget(self.label, self.current_row, 1, 1, 1)
            # create Description box
            self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
            self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
            self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
            self.routing_description_line_edit.setText(description_text)
            # create entry combo
            self.from_load_setEntryBoxSelections(self.current_row,combobox_text.upper())
            # update variables #also adds keyidentifier described above.
            self.row_contents[self.current_row].append(self.label)
            self.row_contents[self.current_row].append(self.routing_description_line_edit)
            self.row_contents[self.current_row].append(self.keyThirdLevelInstruction)
            self.current_row += 1
            self.naming_counter += 1

    def from_load_addAdditionalTechEntryRow(self,combobox_text,label_text):
            # create label
            self.label = QtGui.QLabel(self.centralWidget)
            self.label.setText(label_text)
            self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
            sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
            self.label.setSizePolicy(sizePolicy)
            self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
            # create combobox
            self.from_load_setEntryBoxSelections(self.current_row,combobox_text.upper())
            # update
            self.row_contents[self.current_row].append(self.label)
            self.row_contents[self.current_row].append(self.keyAdditionalTechEntry)
            self.current_row += 1
            self.naming_counter += 1

    def from_load_addRoutingStep(self,combobox_text,label_text,routing_number_text,description_text):
            # reset the routing step numbers so it autopopulates correctly
            self.routing_step_counter = 0
            # create label
            self.label = QtGui.QLabel(self.centralWidget)
            self.label.setText(label_text)
            self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
            sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
            self.label.setSizePolicy(sizePolicy)
            self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
            # create Routing Step Box
            self.routing_step_line_edit = QtGui.QLineEdit(self.centralWidget)
            self.routing_step_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
            self.Addinfogrid.addWidget(self.routing_step_line_edit, self.current_row, 1, 1, 1)
            self.routing_step_line_edit.setFixedWidth(30)
            self.routing_step_line_edit.setText(routing_number_text)
            # create Descr Box
            self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
            self.routing_description_line_edit.setObjectName(
                _fromUtf8("RoutingDescription_" + str(self.naming_counter)))
            self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
            self.routing_description_line_edit.setText(description_text)
            # create entry combo
            self.from_load_setEntryBoxSelections(self.current_row,combobox_text.upper())
            # update variables
            self.row_contents[self.current_row].append(self.label)
            self.row_contents[self.current_row].append(self.routing_step_line_edit)
            self.row_contents[self.current_row].append(self.routing_description_line_edit)
            self.row_contents[self.current_row].append(self.keyRoutingStep)
            self.current_row += 1
            self.naming_counter += 1

    def from_load_add1stLevelInstruction(self, combobox_text, label_text, description_text):
            # create label
            self.label = QtGui.QLabel(self.centralWidget)
            self.label.setText(label_text)
            self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
            sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
            self.label.setSizePolicy(sizePolicy)
            self.Addinfogrid.addWidget(self.label, self.current_row, 0, 1, 1)
            # create Description box
            self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
            self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
            self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
            self.routing_description_line_edit.setText(description_text)
            # create entry combo
            self.from_load_setEntryBoxSelections(self.current_row,combobox_text.upper())
            # update variables #also adds keyidentifier described above.
            self.row_contents[self.current_row].append(self.label)
            self.row_contents[self.current_row].append(self.routing_description_line_edit)
            self.row_contents[self.current_row].append(self.keyFirstLevelInstruction)
            self.current_row += 1
            self.naming_counter += 1

    def from_load_add2ndLevelInstruction(self, combobox_text, label_text, description_text):
            # create label
            self.label = QtGui.QLabel(self.centralWidget)
            self.label.setText(label_text)
            self.label.setObjectName(_fromUtf8("label_" + str(self.naming_counter)))
            sizePolicy = QtGui.QSizePolicy(QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
            self.label.setSizePolicy(sizePolicy)
            self.Addinfogrid.addWidget(self.label, self.current_row, 1, 1, 1)
            # create Description box
            self.routing_description_line_edit = QtGui.QLineEdit(self.centralWidget)
            self.routing_description_line_edit.setObjectName(_fromUtf8("RoutingStep_" + str(self.naming_counter)))
            self.Addinfogrid.addWidget(self.routing_description_line_edit, self.current_row, 2, 1, 1)
            self.routing_description_line_edit.setText(description_text)
            # create entry combo
            self.from_load_setEntryBoxSelections(self.current_row,combobox_text.upper())
            # update variables #also adds keyidentifier described above.
            self.row_contents[self.current_row].append(self.label)
            self.row_contents[self.current_row].append(self.routing_description_line_edit)
            self.row_contents[self.current_row].append(self.keySecondLevelInstruction)
            self.current_row += 1
            self.naming_counter += 1

    def from_load_setEntryBoxSelections(self, current_row,current_selection):
        current_index=0
        if current_selection == 'NONE':
            current_index = 0
        elif current_selection == 'OPERATOR/DATE':
            current_index = 1
        elif current_selection == 'SN/REV':
            current_index = 2
        elif current_selection == 'REV':
            current_index = 3
        elif current_selection == 'ID':
            current_index = 4
        else:
            print("Error in row: "+ str(current_row)+" trying to write " +current_selection)
        combobox = QtGui.QComboBox(self.centralWidget)
        combobox.setObjectName(_fromUtf8("combobox_" + str(self.naming_counter)))
        combobox.addItem(_fromUtf8("None"))
        combobox.addItem(_fromUtf8("Operator/Date"))
        combobox.addItem(_fromUtf8("SN/Rev"))
        combobox.addItem(_fromUtf8("Rev"))
        combobox.addItem(_fromUtf8("ID"))
        self.Addinfogrid.addWidget(combobox, current_row, 3, 1, 1)
        combobox.setCurrentIndex(current_index)
        self.row_contents[self.current_row].append(combobox)

    # used to add more items to comboboxes.
    def setEntryBoxSelections(self, current_row):

        combobox = QtGui.QComboBox(self.centralWidget)
        combobox.setObjectName(_fromUtf8("combobox_" + str(self.naming_counter)))
        combobox.addItem(_fromUtf8("None"))
        combobox.addItem(_fromUtf8("Operator/Date"))
        combobox.addItem(_fromUtf8("SN/Rev"))
        combobox.addItem(_fromUtf8("Rev"))
        combobox.addItem(_fromUtf8("ID"))
        self.Addinfogrid.addWidget(combobox, current_row, 3, 1, 1)
        self.row_contents[self.current_row].append(combobox)

    def retranslateUi(self, BuildMeATraveler): #this function was written by the QT Creator Program
        BuildMeATraveler.setWindowTitle(_translate("BuildMeATraveler", "BuildMeATravler", None))
        self.btnExitTraveler.setText(_translate("BuildMeATraveler", "Exit Traveler", None))
        self.btnSaveTraveler.setText(_translate("BuildMeATraveler", "Save Traveler", None))
        self.btnAddRoutingStep.setText(_translate("BuildMeATraveler", "Add Routing Step", None))
        self.btnAddRoutingInstruction.setText(_translate("BuildMeATraveler", "Add 1st Level Instruction", None))
        self.btnAddSubInstruction.setText(_translate("BuildMeATraveler", "Add 2nd Level Instruction", None))
        self.btnAdd3rdLevelInstruction.setText(_translate("BuildMeATraveler", "Add 3rd Level Instruction", None))
        self.labelPartNumber.setText(_translate("BuildMeATraveler", "Part Number", None))
        self.entryPartDescription.setText(_translate("BuildMeATraveler", "Does Not Work", None))
        self.labelPartDescription.setText(_translate("BuildMeATraveler", "Part Description", None))
        self.btnRemovePreviousEntry.setText(_translate("BuildMeATraveler", "Remove Previous Entry", None))
        self.labelRev.setText(_translate("BuildMeATraveler", "Rev", None))
        self.btnAddTechEntryLine.setText(_translate("BuildMeATraveler", "Add Tech Entry Line", None))
        self.btnSetTemplateLocation.setText(_translate("BuildMeATraveler", "Set Empty Template Location", None))
        self.dropRoutingStepIncrement.setItemText(0, _translate("BuildMeATraveler", "5", None))
        self.dropRoutingStepIncrement.setItemText(1, _translate("BuildMeATraveler", "10", None))
        self.btnPopulateRoutings.setText(_translate("BuildMeATraveler", "Populate Routings", None))
        self.btnClearRoutings.setText(_translate("BuildMeATraveler", "Clear Routings", None))
        self.pushButton.setText(_translate("BuildMeATraveler", "Simple Routing Instruction", None))
        self.menuBuild_A_Traveler.setTitle(_translate("BuildMeATraveler", "Build A Traveler", None))
        self.actionBuildFromTextFile.setText(_translate("BuildMeATraveler", "Build from .txt file", None))


def empty_template_open(self):
    """Allows user to select new empty template location

    :param self: I don't know why i did this. But if it ain't broke, don't fix it.
    :return: None
    """
    try:
        name = QtGui.QFileDialog.getOpenFileName(None, 'Open Empty Traveler Template', 'Save Traveler')
        barcode_template_location = name
        ui.new_empty_template_location = barcode_template_location
    except PackageNotFoundError:
        None

if __name__ == "__main__": #this if statement was written by the QT Creator
    import sys
    app = QtGui.QApplication(sys.argv)
    BuildMeATraveler = QtGui.QMainWindow()
    ui = Ui_BuildMeATraveler()
    ui.setupUi(BuildMeATraveler)
    BuildMeATraveler.show()
    sys.exit(app.exec_())
#end of dynamic methods






